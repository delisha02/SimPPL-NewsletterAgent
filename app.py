import streamlit as st
import requests
import os
import json
from urllib.parse import urlparse
import time
from typing import List, Dict, Optional
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO

# --- API Keys ---
PERPLEXITY_API_KEY = st.secrets.get("PERPLEXITY_API_KEY")

# --- Enhanced Perplexity Fetch with Error Handling ---
def fetch_with_perplexity(query: str, model: str = "sonar-pro") -> str:
    """Enhanced Perplexity API call with better error handling and retry logic."""
    url = "https://api.perplexity.ai/chat/completions"
    headers = {
        "Authorization": f"Bearer {PERPLEXITY_API_KEY}",
        "Content-Type": "application/json"
    }
    payload = {
        "model": model,
        "messages": [{"role": "user", "content": query}],
        "search": True,
        "temperature": 0.2  # Lower temperature for more consistent output
    }
    
    max_retries = 3
    for attempt in range(max_retries):
        try:
            res = requests.post(url, headers=headers, json=payload, timeout=60)
            res.raise_for_status()
            content = res.json()["choices"][0]["message"]["content"]
            # Clean up common LLM artifacts
            content = clean_llm_artifacts(content)
            return content
        except requests.exceptions.Timeout:
            if attempt < max_retries - 1:
                time.sleep(2 ** attempt)  # Exponential backoff
                continue
            return "⏰ Request timed out. Please try again."
        except requests.exceptions.RequestException as e:
            return f"🌐 Network error: {str(e)}"
        except Exception as e:
            return f"❌ Unexpected error: {str(e)}"

def clean_llm_artifacts(text: str) -> str:
    """Remove common LLM artifacts and ensure clean, direct summaries."""
    import re
    
    # Remove citation markers like [1], [2], etc.
    text = re.sub(r'\[\d+\]', '', text)
    
    # Remove common LLM phrases
    llm_phrases = [
        r'certainly[,.]?\s*',
        r'here\'s?\s+(?:a\s+)?(?:summary|breakdown|analysis)[:\s]*',
        r'this\s+(?:article|post|repo|content)\s+(?:discusses|covers|presents)[:\s]*',
        r'in\s+summary[,:]?\s*',
        r'to\s+summarize[,:]?\s*',
        r'based\s+on\s+(?:the\s+)?(?:content|information)[,:]?\s*',
        r'according\s+to\s+(?:the\s+)?(?:article|post|source)[,:]?\s*',
        r'the\s+(?:article|post|content)\s+(?:explains|states|mentions)\s+that\s*',
        r'it\'s\s+worth\s+noting\s+that\s*',
        r'importantly[,:]?\s*',
        r'essentially[,:]?\s*',
        r'basically[,:]?\s*'
    ]
    
    for phrase in llm_phrases:
        text = re.sub(phrase, '', text, flags=re.IGNORECASE)
    
    # Clean up extra whitespace and line breaks
    text = re.sub(r'\n\s*\n', '\n', text)  # Remove double line breaks
    text = re.sub(r'^\s+|\s+$', '', text)  # Trim whitespace
    text = re.sub(r'\s+', ' ', text)  # Normalize spaces
    
    return text

# --- Content Type Detection ---
def detect_content_type(url: str) -> str:
    """Detect the type of content from URL patterns."""
    url_lower = url.lower()
    domain = urlparse(url).netloc.lower()
    
    if 'github.com' in domain:
        if '/releases/' in url or '/tags/' in url:
            return 'github_release'
        elif '/issues/' in url or '/pull/' in url:
            return 'github_issue_pr'
        else:
            return 'github_repo'
    elif 'linkedin.com' in domain:
        return 'linkedin_post'
    elif any(x in domain for x in ['twitter.com', 'x.com']):
        return 'twitter_post'
    elif any(x in domain for x in ['arxiv.org', 'papers.', 'acm.org', 'ieee.org']):
        return 'research_paper'
    elif any(x in domain for x in ['youtube.com', 'youtu.be']):
        return 'video'
    elif any(x in domain for x in ['medium.com', 'substack.com', 'blog']):
        return 'blog_article'
    elif any(x in domain for x in ['docs.', 'documentation']):
        return 'documentation'
    else:
        return 'general_web'

# --- Specialized Summarization Prompts ---
def get_specialized_prompt(url: str, content_type: str) -> str:
    """Generate specialized prompts based on content type with strict output requirements."""
    
    base_instruction = """Write a direct, factual summary in exactly 4-5 lines. Use simple, declarative sentences. Do not include citations [1], [2], etc. Do not use phrases like "certainly", "here's a summary", "this article discusses", or other introductory language. Start immediately with the key information."""
    
    prompts = {
        'github_repo': f"""
{base_instruction}

Analyze this GitHub repository: {url}

Write 4-5 lines covering: the specific problem it solves, key technical features, current development status, and practical use cases. Focus on what developers can accomplish with this tool.
""",
        
        'github_release': f"""
{base_instruction}

Analyze this GitHub release: {url}

Write 4-5 lines covering: major new features added, breaking changes or improvements, performance enhancements, and impact on existing users. Focus on what's changed and why users should update.
""",
        
        'research_paper': f"""
{base_instruction}

Analyze this research paper: {url}

Write 4-5 lines covering: the research problem addressed, methodology used, key findings or contributions, and practical applications. Focus on the scientific contribution and real-world relevance.
""",
        
        'linkedin_post': f"""
{base_instruction}

Analyze this LinkedIn post: {url}

Write 4-5 lines covering: the main professional insight shared, industry context or trend discussed, practical advice or experience, and key takeaway for professionals. Focus on actionable business intelligence.
""",
        
        'twitter_post': f"""
{base_instruction}

Analyze this Twitter post: {url}

Write 4-5 lines covering: the core message or announcement, relevant context or background, industry significance, and community reaction or implications. Focus on the newsworthy information.
""",
        
        'blog_article': f"""
{base_instruction}

Analyze this blog article: {url}

Write 4-5 lines covering: the main argument or thesis, key insights or data presented, practical applications or examples, and implications for readers. Focus on the valuable information and lessons.
""",
        
        'video': f"""
{base_instruction}

Analyze this video: {url}

Write 4-5 lines covering: the main topics covered, key demonstrations or explanations, speaker expertise or unique perspective, and practical value for viewers. Focus on the educational content and takeaways.
""",
        
        'documentation': f"""
{base_instruction}

Analyze this documentation: {url}

Write 4-5 lines covering: the technology or system documented, key features or capabilities explained, target audience or use cases, and practical implementation guidance. Focus on what users can learn and implement.
""",
        
        'general_web': f"""
{base_instruction}

Analyze this web content: {url}

Write 4-5 lines covering: the main subject and key information, relevance to technology or business, notable insights or data, and practical applications. Focus on the essential facts and their significance.
"""
    }
    
    return prompts.get(content_type, prompts['general_web'])

# --- Batch Processing with Progress ---
def process_urls_batch(urls: List[str], model: str = "sonar-pro") -> List[Dict]:
    """Process multiple URLs with progress tracking and error handling."""
    results = []
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    for i, url in enumerate(urls):
        status_text.text(f"Processing {i+1}/{len(urls)}: {url}")
        
        try:
            content_type = detect_content_type(url)
            prompt = get_specialized_prompt(url, content_type)
            summary = fetch_with_perplexity(prompt, model)
            
            # Additional cleaning pass
            summary = post_process_summary(summary)
            
            results.append({
                'url': url,
                'content_type': content_type,
                'summary': summary,
                'status': 'success'
            })
        except Exception as e:
            results.append({
                'url': url,
                'content_type': 'unknown',
                'summary': f"Failed to process: {str(e)}",
                'status': 'error'
            })
        
        progress_bar.progress((i + 1) / len(urls))
        time.sleep(1)  # Rate limiting - increased to be more conservative
    
    status_text.empty()
    progress_bar.empty()
    return results

def post_process_summary(summary: str) -> str:
    """Final cleanup pass to ensure quality output."""
    import re
    
    # Split into sentences and clean each one
    sentences = summary.split('.')
    clean_sentences = []
    
    for sentence in sentences:
        sentence = sentence.strip()
        if len(sentence) < 10:  # Skip very short fragments
            continue
            
        # Remove remaining artifacts
        sentence = re.sub(r'^(the\s+)?content\s+(shows|reveals|indicates)\s+that\s+', '', sentence, flags=re.IGNORECASE)
        sentence = re.sub(r'^(this\s+)?(article|post|repo|video)\s+(is|focuses|covers)\s+', '', sentence, flags=re.IGNORECASE)
        
        if sentence:
            clean_sentences.append(sentence)
    
    # Rejoin and ensure proper formatting
    result = '. '.join(clean_sentences)
    if result and not result.endswith('.'):
        result += '.'
    
    return result

# --- Export Functionality ---
def export_summaries(results: List[Dict], format_type: str = "markdown"):
    """Export summaries in various formats."""
    if format_type == "markdown":
        content = "# Newsletter Link Summaries\n\n"
        for result in results:
            if result['status'] == 'success':
                content += f"## 🔗 [{result['url']}]({result['url']})\n"
                content += f"**Type:** {result['content_type'].replace('_', ' ').title()}\n\n"
                content += f"{result['summary']}\n\n---\n\n"
        return content
    elif format_type == "json":
        return json.dumps(results, indent=2)
    elif format_type == "docx":
        return create_word_document(results)

def create_word_document(results: List[Dict]) -> BytesIO:
    """Create a Word document with the summaries."""
    doc = Document()
    
    # Add title
    title = doc.add_heading('Newsletter Link Summaries', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add creation date
    from datetime import datetime
    date_para = doc.add_paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add summary statistics
    total_urls = len(results)
    successful = sum(1 for r in results if r['status'] == 'success')
    
    stats_para = doc.add_paragraph()
    stats_para.add_run(f"Total URLs Analyzed: {total_urls} | Successful: {successful}").bold = True
    stats_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()  # Add space
    
    # Add individual summaries
    for i, result in enumerate(results, 1):
        if result['status'] == 'success':
            # Add URL as heading
            url_heading = doc.add_heading(f"{i}. Link Analysis", level=1)
            
            # Add URL
            url_para = doc.add_paragraph()
            url_para.add_run("URL: ").bold = True
            url_para.add_run(result['url'])
            
            # Add content type
            type_para = doc.add_paragraph()
            type_para.add_run("Content Type: ").bold = True
            type_para.add_run(result['content_type'].replace('_', ' ').title())
            
            # Add summary
            summary_para = doc.add_paragraph()
            summary_para.add_run("Summary:").bold = True
            doc.add_paragraph(result['summary'])
            
            # Add separator
            doc.add_paragraph("_" * 50)
            doc.add_paragraph()  # Add space
    
    # Save to BytesIO
    doc_buffer = BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer

# --- Streamlit UI ---
st.set_page_config(
    page_title="🔗 Enhanced Newsletter Summarizer", 
    layout="wide",
    initial_sidebar_state="expanded"
)

st.title("🔗 Enhanced Newsletter Link Summarizer")
st.markdown("AI-powered content analysis for tech newsletters with clean, direct summaries.")

# --- Sidebar Configuration ---
with st.sidebar:
    st.markdown("### ⚙️ Configuration")
    
    # Model selection
    model_choice = st.selectbox(
        "Perplexity Model:",
        ["sonar-pro", "sonar"],
        help="Choose the Perplexity model for analysis"
    )
    
    # Export options
    st.markdown("### 📤 Export Options")
    export_format = st.radio(
        "Export Format:",
        ["markdown", "json", "docx"]
    )
    
    # Content type info
    st.markdown("### 📋 Supported Content Types")
    st.markdown("""
    - 🔧 GitHub repositories & releases
    - 📄 Research papers (arXiv, ACM, IEEE)
    - 💼 LinkedIn posts
    - 🐦 Twitter/X posts
    - 📝 Blog articles & Medium posts
    - 🎥 YouTube videos
    - 📚 Documentation
    - 🌐 General web content
    """)
    
    st.markdown("### ✨ Output Features")
    st.markdown("""
    - Clean, direct 4-5 line summaries
    - No citation markers or LLM artifacts
    - Content-specific analysis
    - Professional newsletter format
    """)

# --- Main Interface ---
col1, col2 = st.columns([2, 1])

with col1:
    st.markdown("### 📝 Enter URLs")
    urls_input = st.text_area(
        "Paste URLs (one per line):",
        height=300,
        placeholder="""https://github.com/user/repo
https://arxiv.org/abs/2301.00000
https://www.linkedin.com/posts/user-post
https://twitter.com/user/status/123
https://medium.com/@user/article"""
    )

with col2:
    st.markdown("### 🎯 Quick Actions")
    
    if st.button("🧠 Analyze Links", type="primary"):
        urls = [url.strip() for url in urls_input.splitlines() if url.strip()]
        
        if not urls:
            st.warning("Please enter at least one URL.")
        elif not PERPLEXITY_API_KEY:
            st.error("Perplexity API key not found. Please check your secrets configuration.")
        else:
            with st.spinner("Processing URLs..."):
                st.session_state.results = process_urls_batch(urls, model_choice)
    
    if st.button("🧹 Clear Results"):
        if 'results' in st.session_state:
            del st.session_state.results
        st.rerun()

# --- Results Display ---
if 'results' in st.session_state and st.session_state.results:
    st.markdown("---")
    st.markdown("## 📊 Analysis Results")
    
    # Summary stats
    total_urls = len(st.session_state.results)
    successful = sum(1 for r in st.session_state.results if r['status'] == 'success')
    failed = total_urls - successful
    
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("Total URLs", total_urls)
    with col2:
        st.metric("Successful", successful, delta=None if failed == 0 else f"-{failed} failed")
    with col3:
        st.metric("Success Rate", f"{(successful/total_urls)*100:.1f}%")
    
    # Export button
    if successful > 0:
        if export_format == "docx":
            doc_buffer = export_summaries(st.session_state.results, export_format)
            st.download_button(
                label="📄 Export as Word Document",
                data=doc_buffer.getvalue(),
                file_name="newsletter_summaries.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            export_content = export_summaries(st.session_state.results, export_format)
            st.download_button(
                label=f"📤 Export as {export_format.upper()}",
                data=export_content,
                file_name=f"newsletter_summaries.{export_format}",
                mime="text/markdown" if export_format == "markdown" else "application/json"
            )
    
    # Individual results
    for i, result in enumerate(st.session_state.results):
        if result['status'] == 'success':
            with st.expander(f"🔗 {result['url']}", expanded=True):
                col1, col2 = st.columns([3, 1])
                with col1:
                    st.markdown(f"**Content Type:** {result['content_type'].replace('_', ' ').title()}")
                with col2:
                    st.markdown(f"**Status:** ✅ Success")
                
                st.markdown("**Summary:**")
                st.info(result['summary'])
        else:
            with st.expander(f"❌ {result['url']}", expanded=False):
                st.error(f"Failed to process: {result['summary']}")

# --- Footer ---
st.markdown("---")
st.markdown("""
<div style='text-align: center; color: #666; font-size: 14px;'>
    Enhanced Newsletter Summarizer | Powered by Perplexity AI | Clean Output Guaranteed
</div>
""", unsafe_allow_html=True)