import streamlit as st
import requests
import os
import json
from urllib.parse import urlparse
import time
from typing import List, Dict, Optional
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO
from bs4 import BeautifulSoup
import base64
from mistralai import Mistral
import re

MISTRAL_API_KEY = st.secrets.get("MISTRAL_API_KEY")
if MISTRAL_API_KEY:
    mistral_client = Mistral(api_key=MISTRAL_API_KEY)
else:
    mistral_client = None

# --- API Keys ---
PERPLEXITY_API_KEY = st.secrets.get("PERPLEXITY_API_KEY")

# --- FIXED: Direct Content Extraction ---
def extract_content_from_url(url: str) -> Dict[str, str]:
    """Extract actual content from URL using web scraping."""
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8",
        "Accept-Language": "en-US,en;q=0.5",
        "Accept-Encoding": "gzip, deflate",
        "Connection": "keep-alive",
        "Upgrade-Insecure-Requests": "1"
    }
    
    try:
        response = requests.get(url, headers=headers, timeout=15)
        response.raise_for_status()
        soup = BeautifulSoup(response.content, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style", "nav", "footer", "header"]):
            script.decompose()
        
        # Extract different types of content based on URL
        content_type = detect_content_type(url)
        
        if content_type == 'twitter_post':
            return extract_twitter_content(soup, url)
        elif content_type == 'github_repo':
            return extract_github_content(soup, url)
        elif content_type == 'linkedin_post':
            return extract_linkedin_content(soup, url)
        else:
            return extract_general_content(soup, url)
            
    except requests.RequestException as e:
        return {
            "content": f"Failed to fetch content: {str(e)}",
            "title": "Error",
            "success": False
        }

# Quick fix: Replace your extract_twitter_content function with this improved version

def extract_twitter_content(soup: BeautifulSoup, url: str) -> Dict[str, str]:
    """Enhanced Twitter/X post content extraction with Perplexity fallback."""
    
    # First try traditional scraping methods
    tweet_text = ""
    
    # Try multiple meta tag approaches
    meta_selectors = [
        ('meta', {'property': 'og:description'}),
        ('meta', {'name': 'twitter:description'}),
        ('meta', {'name': 'description'}),
        ('meta', {'property': 'og:title'}),
        ('meta', {'name': 'twitter:title'})
    ]
    
    for tag, attrs in meta_selectors:
        meta_tag = soup.find(tag, attrs)
        if meta_tag and meta_tag.get('content'):
            content = meta_tag.get('content', '').strip()
            if content and len(content) > 20:  # Ensure we have substantial content
                tweet_text = content
                break
    
    # Clean up common meta tag artifacts
    if tweet_text:
        # Remove quotes and common prefixes
        tweet_text = tweet_text.strip('"').strip("'")
        prefixes_to_remove = [
            'See new posts', 'Latest tweets from', 'Tweets by',
            'Check out the latest tweets from'
        ]
        for prefix in prefixes_to_remove:
            if tweet_text.startswith(prefix):
                tweet_text = tweet_text[len(prefix):].strip()
    
    # Get title/author info
    title = "Twitter Post"
    title_tag = soup.find('title')
    if title_tag:
        title_text = title_tag.get_text()
        if ' on X:' in title_text or ' on Twitter:' in title_text:
            author = title_text.split(' on ')[0].strip()
            title = f"Twitter Post by {author}"
    
    # If traditional scraping fails, use Perplexity as fallback
    if not tweet_text or len(tweet_text) < 20:
        st.info(f"Traditional scraping failed for {url}, using AI analysis...")
        tweet_text = analyze_twitter_url_with_perplexity(url)
        title = "Twitter Post (AI Analysis)"
        
        return {
            "content": tweet_text,
            "title": title,
            "success": True  # Perplexity analysis counts as success
        }
    
    return {
        "content": tweet_text,
        "title": title,
        "success": True
    }

def analyze_twitter_url_with_perplexity(url: str) -> str:
    """Use Perplexity to analyze Twitter URL when scraping fails."""
    prompt = f"""
Analyze this Twitter/X post URL and provide a concise summary: {url}

Focus on:
- Main message or content of the tweet
- Who posted it (if identifiable from URL or public info)
- Key points, insights, or announcements
- Any relevant context

Provide a direct, factual summary in 3-4 sentences without introductory phrases.
"""
    
    try:
        return fetch_with_perplexity_direct(prompt)
    except Exception as e:
        return f"Could not analyze Twitter post: {str(e)}"

def extract_github_content(soup: BeautifulSoup, url: str) -> Dict[str, str]:
    """Extract GitHub repository content."""
    # Repository description
    description = ""
    desc_elem = soup.find('p', class_='f4 my-3') or soup.find('[data-testid="repo-description"]')
    if desc_elem:
        description = desc_elem.get_text(strip=True)
    
    # README content (first few paragraphs)
    readme_content = ""
    readme_elem = soup.find('div', {'data-testid': 'readme'}) or soup.find('article', class_='markdown-body')
    if readme_elem:
        paragraphs = readme_elem.find_all(['p', 'h1', 'h2', 'h3'], limit=5)
        readme_content = " ".join([p.get_text(strip=True) for p in paragraphs])
    
    # Repository stats
    stats = ""
    star_elem = soup.find('a', href=lambda x: x and '/stargazers' in x)
    if star_elem:
        stars = star_elem.get_text(strip=True)
        stats = f"Stars: {stars}"
    
    title = soup.find('title').get_text() if soup.find('title') else "GitHub Repository"
    
    combined_content = f"{description} {readme_content} {stats}".strip()
    
    return {
        "content": combined_content if combined_content else "Could not extract repository content",
        "title": title,
        "success": bool(combined_content)
    }

# Quick fix: Replace your extract_linkedin_content function with this

def extract_linkedin_content(soup: BeautifulSoup, url: str) -> Dict[str, str]:
    """Enhanced LinkedIn post content extraction with AI fallback."""
    
    # Try enhanced meta tag extraction first
    post_content = ""
    author = ""
    
    # Enhanced meta tag selectors
    meta_selectors = [
        ('meta', {'property': 'og:description'}),
        ('meta', {'property': 'og:title'}),
        ('meta', {'name': 'twitter:description'}),
        ('meta', {'name': 'description'}),
        ('meta', {'name': 'title'})
    ]
    
    for tag_name, attrs in meta_selectors:
        meta_tag = soup.find(tag_name, attrs)
        if meta_tag and meta_tag.get('content'):
            content = meta_tag.get('content', '').strip()
            # Filter out generic LinkedIn content
            if (content and 
                len(content) > 30 and 
                'Join LinkedIn' not in content and
                'Sign up to see' not in content and
                'LinkedIn' != content):
                post_content = content
                break
    
    # Extract author from URL
    if '/posts/' in url:
        try:
            username_part = url.split('/posts/')[1].split('_')[0]
            author = username_part.replace('-', ' ').replace('phd', 'PhD').title()
        except:
            pass
    
    # Clean up content
    if post_content:
        # Remove quotes and LinkedIn artifacts
        post_content = post_content.strip('"\'')
        post_content = re.sub(r'^LinkedIn:?\s*', '', post_content, flags=re.IGNORECASE)
        post_content = re.sub(r'\s*-\s*LinkedIn$', '', post_content, flags=re.IGNORECASE)
    
    # If meta extraction fails or content is too generic, use AI analysis
    if not post_content or len(post_content) < 30:
        ai_content = analyze_linkedin_with_ai(url, author)
        return {
            "content": ai_content,
            "title": f"LinkedIn Post by {author}" if author else "LinkedIn Post (AI Analysis)",
            "success": True
        }
    
    return {
        "content": post_content,
        "title": f"LinkedIn Post by {author}" if author else "LinkedIn Post",
        "success": True
    }

def analyze_linkedin_with_ai(url: str, author: str = "") -> str:
    """Analyze LinkedIn URL using Perplexity when scraping fails."""
    
    author_context = f" by {author}" if author else ""
    
    prompt = f"""
Analyze this LinkedIn post URL and provide a professional summary: {url}

Focus on:
- Main message, insight, or professional advice shared
- Author's expertise or perspective{author_context}
- Key points or takeaways relevant to professionals
- Industry context or significance
- Any specific tools, concepts, or strategies mentioned

Provide a clear, informative summary in 3-4 sentences that would be valuable for a professional newsletter.
Be direct and factual, avoiding introductory phrases.
"""
    
    try:
        return fetch_with_perplexity_direct(prompt)
    except Exception as e:
        return f"Could not analyze LinkedIn post. The post may require authentication to access or may be private. Author{author_context} shared professional insights that are not publicly accessible."
def extract_general_content(soup: BeautifulSoup, url: str) -> Dict[str, str]:
    """Extract general web content."""
    # Try to find main content
    main_content = ""
    
    # Priority selectors for main content
    content_selectors = [
        'main',
        'article',
        '.post-content',
        '.entry-content',
        '.content',
        '[role="main"]',
        '.main-content'
    ]
    
    for selector in content_selectors:
        content_elem = soup.select_one(selector)
        if content_elem:
            # Get first few paragraphs
            paragraphs = content_elem.find_all(['p', 'h1', 'h2', 'h3'], limit=10)
            main_content = " ".join([p.get_text(strip=True) for p in paragraphs if len(p.get_text(strip=True)) > 20])
            break
    
    # Fallback to all paragraphs
    if not main_content:
        paragraphs = soup.find_all('p', limit=10)
        main_content = " ".join([p.get_text(strip=True) for p in paragraphs if len(p.get_text(strip=True)) > 20])
    
    title = "Web Content"
    if soup.find('title'):
        title = soup.find('title').get_text()
    
    return {
        "content": main_content[:2000] if main_content else "Could not extract content",  # Limit content length
        "title": title,
        "success": bool(main_content)
    }

# --- FIXED: Content-Based Summarization ---
def summarize_extracted_content(content: str, url: str, content_type: str) -> str:
    """Summarize the actual extracted content using Perplexity."""
    if not content or content.startswith("Could not extract") or content.startswith("Failed to fetch"):
        return f"❌ Unable to extract content from {url}"
    
    # Create a focused prompt for the extracted content
    prompt = f"""
Analyze this content and write a direct, factual summary in exactly 4-5 lines. Use simple, declarative sentences. Do not include citations, introductory phrases, or meta-commentary. Start immediately with the key information.

Content Type: {content_type}
Source URL: {url}

Content to summarize:
{content}

Focus on: main points, key insights, actionable information, and significance. Write as if for a professional newsletter.
"""
    
    return fetch_with_perplexity_direct(prompt)

def fetch_with_perplexity_direct(prompt: str, model: str = "sonar-pro") -> str:
    """Direct Perplexity API call for content summarization."""
    url = "https://api.perplexity.ai/chat/completions"
    headers = {
        "Authorization": f"Bearer {PERPLEXITY_API_KEY}",
        "Content-Type": "application/json"
    }
    payload = {
        "model": model,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.2,
        "max_tokens": 300  # Limit response length
    }
    
    max_retries = 3
    for attempt in range(max_retries):
        try:
            res = requests.post(url, headers=headers, json=payload, timeout=30)
            res.raise_for_status()
            content = res.json()["choices"][0]["message"]["content"]
            return clean_llm_artifacts(content)
        except requests.exceptions.Timeout:
            if attempt < max_retries - 1:
                time.sleep(2 ** attempt)
                continue
            return "⏰ Request timed out. Please try again."
        except requests.exceptions.RequestException as e:
            return f"🌐 Network error: {str(e)}"
        except Exception as e:
            return f"❌ Unexpected error: {str(e)}"

# --- Enhanced Content Validation ---
def validate_summary_quality(summary: str, original_content: str) -> bool:
    """Validate that summary is actually about the content, not generic info."""
    summary_lower = summary.lower()
    
    # Red flags indicating generic/biographical content instead of post content
    red_flags = [
        "co-founder", "ceo of", "previously worked", "completed her phd",
        "research scientist", "university", "publications", "collaborated with",
        "founded", "works at", "degree from", "author of", "professor at"
    ]
    
    # Check if summary contains too many biographical red flags
    red_flag_count = sum(1 for flag in red_flags if flag in summary_lower)
    
    # If more than 2 red flags, likely biographical content
    if red_flag_count > 2:
        return False
    
    # Check if summary mentions specific content elements
    content_words = set(original_content.lower().split())
    summary_words = set(summary_lower.split())
    
    # Calculate content overlap
    overlap = len(content_words.intersection(summary_words))
    overlap_ratio = overlap / max(len(summary_words), 1)
    
    # Summary should have reasonable overlap with original content
    return overlap_ratio > 0.1

# --- FIXED: Process URLs with Content Extraction ---
def process_urls_batch(urls: List[str], model: str = "sonar-pro") -> List[Dict]:
    """Process multiple URLs with actual content extraction."""
    results = []
    progress_bar = st.progress(0)
    status_text = st.empty()

    for i, url in enumerate(urls):
        status_text.text(f"Processing {i+1}/{len(urls)}: {url}")
        
        try:
            # Step 1: Extract actual content from URL
            extracted = extract_content_from_url(url)
            content_type = detect_content_type(url)
            
            if not extracted['success']:
                results.append({
                    'url': url,
                    'content_type': content_type,
                    'summary': f"❌ {extracted['content']}",
                    'status': 'error'
                })
                continue
            
            # Step 2: Summarize the extracted content
            summary = summarize_extracted_content(extracted['content'], url, content_type)
            
            # Step 3: Validate summary quality
            if not validate_summary_quality(summary, extracted['content']):
                summary = f"⚠️ Generated summary may not accurately reflect the specific post content. Content extraction may have failed for this URL type."
            
            # Step 4: Final cleanup
            summary = post_process_summary(summary)
            
            results.append({
                'url': url,
                'content_type': content_type,
                'summary': summary,
                'status': 'success',
                'extracted_content': extracted['content'][:200] + "..." if len(extracted['content']) > 200 else extracted['content']  # For debugging
            })
            
        except Exception as e:
            results.append({
                'url': url,
                'content_type': 'unknown',
                'summary': f"❌ Failed to process: {str(e)}",
                'status': 'error'
            })

        progress_bar.progress((i + 1) / len(urls))
        time.sleep(1)  # Rate limiting

    status_text.empty()
    progress_bar.empty()
    return results

# --- Keep Original Helper Functions ---
def clean_llm_artifacts(text: str) -> str:
    """Remove common LLM artifacts and ensure clean, direct summaries."""
    # Remove citation markers like [1], [2], etc.
    text = re.sub(r'\[\d+\]', '', text)
    
    # Remove common LLM phrases
    llm_phrases = [
        r'certainly[,.]?\s*',
        r'here\'s?\s+(?:a\s+)?(?:summary|breakdown|analysis)[:\s]*',
        r'this\s+(?:article|post|repo|content)\s+(?:discusses|covers|presents)[:\s]*',
        r'in\s+summary[,:]?\s*',
        r'to\s+summarize[,:]?\s*',
        r'based\s+on\s+(?:the\s+)?(?:content|information)[,:]?\s*',
        r'according\s+to\s+(?:the\s+)?(?:article|post|source)[,:]?\s*',
        r'the\s+(?:article|post|content)\s+(?:explains|states|mentions)\s+that\s*',
        r'it\'s\s+worth\s+noting\s+that\s*',
        r'importantly[,:]?\s*',
        r'essentially[,:]?\s*',
        r'basically[,:]?\s*'
    ]
    
    for phrase in llm_phrases:
        text = re.sub(phrase, '', text, flags=re.IGNORECASE)
    
    # Clean up extra whitespace and line breaks
    text = re.sub(r'\n\s*\n', '\n', text)
    text = re.sub(r'^\s+|\s+$', '', text)
    text = re.sub(r'\s+', ' ', text)
    
    return text

def detect_content_type(url: str) -> str:
    """Detect the type of content from URL patterns."""
    url_lower = url.lower()
    domain = urlparse(url).netloc.lower()
    
    if 'github.com' in domain:
        if '/releases/' in url or '/tags/' in url:
            return 'github_release'
        elif '/issues/' in url or '/pull/' in url:
            return 'github_issue_pr'
        else:
            return 'github_repo'
    elif 'linkedin.com' in domain:
        return 'linkedin_post'
    elif any(x in domain for x in ['twitter.com', 'x.com']):
        return 'twitter_post'
    elif any(x in domain for x in ['arxiv.org', 'papers.', 'acm.org', 'ieee.org']):
        return 'research_paper'
    elif any(x in domain for x in ['youtube.com', 'youtu.be']):
        return 'video'
    elif any(x in domain for x in ['medium.com', 'substack.com', 'blog']):
        return 'blog_article'
    elif any(x in domain for x in ['docs.', 'documentation']):
        return 'documentation'
    else:
        return 'general_web'

def post_process_summary(summary: str) -> str:
    """Final cleanup pass to ensure quality output."""
    # Split into sentences and clean each one
    sentences = summary.split('.')
    clean_sentences = []
    
    for sentence in sentences:
        sentence = sentence.strip()
        if len(sentence) < 10:  # Skip very short fragments
            continue
            
        # Remove remaining artifacts
        sentence = re.sub(r'^(the\s+)?content\s+(shows|reveals|indicates)\s+that\s+', '', sentence, flags=re.IGNORECASE)
        sentence = re.sub(r'^(this\s+)?(article|post|repo|video)\s+(is|focuses|covers)\s+', '', sentence, flags=re.IGNORECASE)
        
        if sentence:
            clean_sentences.append(sentence)
    
    # Rejoin and ensure proper formatting
    result = '. '.join(clean_sentences)
    if result and not result.endswith('.'):
        result += '.'
    
    return result

# --- Keep Original Export Functions ---
def export_summaries(results: List[Dict], format_type: str = "markdown"):
    """Export summaries in various formats."""
    if format_type == "markdown":
        content = "# Newsletter Link Summaries\n\n"
        for result in results:
            if result['status'] == 'success':
                content += f"## 🔗 [{result['url']}]({result['url']})\n"
                content += f"**Type:** {result['content_type'].replace('_', ' ').title()}\n\n"
                content += f"{result['summary']}\n\n---\n\n"
        return content
    elif format_type == "json":
        return json.dumps(results, indent=2)
    elif format_type == "docx":
        return create_word_document(results)

def create_word_document(results: List[Dict]) -> BytesIO:
    """Create a Word document with the summaries."""
    doc = Document()
    
    # Add title
    title = doc.add_heading('Newsletter Link Summaries', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add creation date
    from datetime import datetime
    date_para = doc.add_paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add summary statistics
    total_urls = len(results)
    successful = sum(1 for r in results if r['status'] == 'success')
    
    stats_para = doc.add_paragraph()
    stats_para.add_run(f"Total URLs Analyzed: {total_urls} | Successful: {successful}").bold = True
    stats_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()  # Add space
    
    # Add individual summaries
    for i, result in enumerate(results, 1):
        if result['status'] == 'success':
            # Add URL as heading
            url_heading = doc.add_heading(f"{i}. Link Analysis", level=1)
            
            # Add URL
            url_para = doc.add_paragraph()
            url_para.add_run("URL: ").bold = True
            url_para.add_run(result['url'])
            
            # Add content type
            type_para = doc.add_paragraph()
            type_para.add_run("Content Type: ").bold = True
            type_para.add_run(result['content_type'].replace('_', ' ').title())
            
            # Add summary
            summary_para = doc.add_paragraph()
            summary_para.add_run("Summary:").bold = True
            doc.add_paragraph(result['summary'])
            
            # Add separator
            doc.add_paragraph("_" * 50)
            doc.add_paragraph()  # Add space
    
    # Save to BytesIO
    doc_buffer = BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer

# --- Streamlit UI (Keep Original) ---
st.set_page_config(
    page_title="🔗 Enhanced Newsletter Summarizer", 
    layout="wide",
    initial_sidebar_state="expanded"
)

st.title("🔗 Enhanced Newsletter Link Summarizer")
st.markdown("AI-powered content analysis for tech newsletters with clean, direct summaries.")

# --- Initialize Session State ---
if "model_choice" not in st.session_state:
    st.session_state.model_choice = "sonar-pro"

# --- Sidebar Configuration ---
with st.sidebar:
    st.markdown("### ⚙️ Configuration")
    
    # Model selection
    model_choice = st.selectbox(
        "Perplexity Model:",
        ["sonar-pro", "sonar"],
        index=0 if st.session_state.model_choice == "sonar-pro" else 1
    )
    
    # Update session state
    st.session_state.model_choice = model_choice
    
    # Export options
    st.markdown("### 📤 Export Options")
    export_format = st.radio(
        "Export Format:",
        ["markdown", "json", "docx"]
    )
    
    # Content type info
    st.markdown("### 📋 Supported Content Types")
    st.markdown("""
    - 🔧 GitHub repositories & releases
    - 📄 Research papers (arXiv, ACM, IEEE)
    - 💼 LinkedIn posts
    - 🐦 Twitter/X posts
    - 📝 Blog articles & Medium posts
    - 🎥 YouTube videos
    - 📚 Documentation
    - 🌐 General web content
    """)
    
    st.markdown("### ✨ New Features")
    st.markdown("""
    - **Direct content extraction** from URLs
    - **Content validation** to ensure accuracy
    - **Improved Twitter/X post handling**
    - **Better error handling** and debugging
    """)
    
    # API Status
    st.markdown("### 🔑 API Status")
    if PERPLEXITY_API_KEY:
        st.success("✅ Perplexity API: Connected")
    else:
        st.error("❌ Perplexity API: Not configured")
        
    if MISTRAL_API_KEY:
        st.success("✅ Mistral API: Connected")
    else:
        st.warning("⚠️ Mistral API: Not configured (OCR disabled)")

# --- Main Interface ---
col1, col2 = st.columns([2, 1])

with col1:
    st.markdown("### 📝 Enter URLs")
    urls_input = st.text_area(
        "Paste URLs (one per line):",
        height=300,
        placeholder="""https://github.com/user/repo
https://arxiv.org/abs/2301.00000
https://www.linkedin.com/posts/user-post
https://twitter.com/user/status/123
https://medium.com/@user/article"""
    )

with col2:
    st.markdown("### 🎯 Quick Actions")
    
    if st.button("🧠 Analyze Links", type="primary"):
        urls = [url.strip() for url in urls_input.splitlines() if url.strip()]
        
        if not urls:
            st.warning("Please enter at least one URL.")
        elif not PERPLEXITY_API_KEY:
            st.error("Perplexity API key not found. Please check your secrets configuration.")
        else:
            with st.spinner("Extracting content and generating summaries..."):
                st.session_state.results = process_urls_batch(urls, model_choice)
    
    if st.button("🧹 Clear Results"):
        if 'results' in st.session_state:
            del st.session_state.results
        st.rerun()

# --- Results Display ---
if 'results' in st.session_state and st.session_state.results:
    st.markdown("---")
    st.markdown("## 📊 Analysis Results")
    
    # Summary stats
    total_urls = len(st.session_state.results)
    successful = sum(1 for r in st.session_state.results if r['status'] == 'success')
    failed = total_urls - successful
    
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("Total URLs", total_urls)
    with col2:
        st.metric("Successful", successful, delta=None if failed == 0 else f"-{failed} failed")
    with col3:
        st.metric("Success Rate", f"{(successful/total_urls)*100:.1f}%")
    
    # Export button
    if successful > 0:
        if export_format == "docx":
            doc_buffer = export_summaries(st.session_state.results, export_format)
            st.download_button(
                label="📄 Export as Word Document",
                data=doc_buffer.getvalue(),
                file_name="newsletter_summaries.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            export_content = export_summaries(st.session_state.results, export_format)
            st.download_button(
                label=f"📤 Export as {export_format.upper()}",
                data=export_content,
                file_name=f"newsletter_summaries.{export_format}",
                mime="text/markdown" if export_format == "markdown" else "application/json"
            )
    
    # Individual results with debugging info
    for i, result in enumerate(st.session_state.results):
        if result['status'] == 'success':
            with st.expander(f"🔗 {result['url']}", expanded=True):
                col1, col2 = st.columns([3, 1])
                with col1:
                    st.markdown(f"**Content Type:** {result['content_type'].replace('_', ' ').title()}")
                with col2:
                    st.markdown(f"**Status:** ✅ Success")
                
                st.markdown("**Summary:**")
                st.info(result['summary'])
                
                # Debug info (optional)
                if st.checkbox(f"Show extracted content preview", key=f"debug_{i}"):
                    if 'extracted_content' in result:
                        st.markdown("**Extracted Content Preview:**")
                        st.code(result['extracted_content'], language="text")
        else:
            with st.expander(f"❌ {result['url']}", expanded=False):
                st.error(f"Failed to process: {result['summary']}")

# --- Footer ---
st.markdown("---")
st.markdown("""
<div style='text-align: center; color: #666; font-size: 14px;'>
    Enhanced Newsletter Summarizer v2.0 | Direct Content Extraction | Powered by Perplexity AI
</div>
""", unsafe_allow_html=True)